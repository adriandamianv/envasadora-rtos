# -*- coding: utf-8 -*-
"""Genera Informe_Final.docx con python-docx a partir de contenido.py.
El índice se inserta como campo TOC real: Word lo llena al abrir
(clic derecho → Actualizar campo, o Ctrl+A y F9)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

from contenido import BLOQUES, INTEGRANTES, PORTADA

AQUI = Path(__file__).parent


def campo_toc(parrafo):
    """Inserta el campo TOC de Word (índice automático de verdad)."""
    run = parrafo.add_run()
    inicio = OxmlElement("w:fldChar")
    inicio.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    separador = OxmlElement("w:fldChar")
    separador.set(qn("w:fldCharType"), "separate")
    aviso = OxmlElement("w:t")
    aviso.text = "Índice: clic derecho → Actualizar campo."
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    run._r.append(inicio)
    run._r.append(instr)
    run._r.append(separador)
    run._r.append(aviso)
    run._r.append(fin)


def centrado(doc, texto, tam, negrita=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(texto)
    r.font.size = Pt(tam)
    r.bold = negrita
    return p


def main():
    doc = Document()
    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)

    # ---- portada ----
    # logos: UG a la izquierda, Facultad Industrial a la derecha
    t_logos = doc.add_table(rows=1, cols=3)
    t_logos.autofit = True
    c_izq, _, c_der = t_logos.rows[0].cells
    p_izq = c_izq.paragraphs[0]
    p_izq.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_izq.add_run().add_picture(str(AQUI / PORTADA["logo_izq"]), width=Inches(1.35))
    p_der = c_der.paragraphs[0]
    p_der.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_der.add_run().add_picture(str(AQUI / PORTADA["logo_der"]), width=Inches(1.1))
    doc.add_paragraph()
    centrado(doc, PORTADA["universidad"], 16, True)
    centrado(doc, PORTADA["carrera"], 12)
    doc.add_paragraph()
    centrado(doc, PORTADA["titulo"], 15, True)
    doc.add_paragraph()
    centrado(doc, f"Materia: {PORTADA['materia']}", 12)
    centrado(doc, f"Docente: {PORTADA['profesor']}", 12)
    centrado(doc, PORTADA["periodo"], 12)
    doc.add_paragraph()
    etiqueta = "Integrante:" if len(INTEGRANTES) == 1 else "Integrantes:"
    centrado(doc, etiqueta, 12, True)
    for nombre in INTEGRANTES:
        centrado(doc, nombre, 12)
    doc.add_paragraph()
    centrado(doc, PORTADA["fecha"], 12)
    doc.add_page_break()

    # ---- índice automático ----
    doc.add_heading("Índice", level=1)
    campo_toc(doc.add_paragraph())
    doc.add_page_break()

    # ---- cuerpo ----
    for bloque in BLOQUES:
        tipo = bloque[0]
        if tipo == "h1":
            doc.add_heading(bloque[1], level=1)
        elif tipo == "h2":
            doc.add_heading(bloque[1], level=2)
        elif tipo == "p":
            doc.add_paragraph(bloque[1])
        elif tipo == "lista":
            for item in bloque[1]:
                doc.add_paragraph(item, style="List Bullet")
        elif tipo == "tabla":
            _, cab, filas = bloque
            t = doc.add_table(rows=1, cols=len(cab))
            t.style = "Light Grid Accent 1"
            for i, c in enumerate(cab):
                t.rows[0].cells[i].text = c
            for fila in filas:
                celdas = t.add_row().cells
                for i, v in enumerate(fila):
                    celdas[i].text = v
            doc.add_paragraph()
        elif tipo == "img":
            _, ruta, pie = bloque
            ancho = Inches(6.0 if "lcd" not in ruta else 3.2)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(AQUI / ruta), width=ancho)
            pie_p = doc.add_paragraph()
            pie_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = pie_p.add_run(pie)
            r.italic = True
            r.font.size = Pt(9)

    salida = AQUI / "Informe_Final.docx"
    doc.save(str(salida))
    print(f"OK: {salida} ({salida.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
